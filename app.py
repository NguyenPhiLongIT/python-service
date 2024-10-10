import os 
from flask import Flask, request, jsonify, Response
from flask_cors import CORS
from neo4j import GraphDatabase
import marko
import google.generativeai as genai
import PyPDF2
from docx import Document

app = Flask(__name__)
CORS(app)

URI = "neo4j+s://5ef0fe4c.databases.neo4j.io"
AUTH = ("neo4j", os.getenv("PASSWORD_NEO4J"))
dict = {
    "Concept": "Định nghĩa",
    "Example": "Ví dụ",
    "EnglishName": "Tên tiếng anh",
    "OtherName": "Tên gọi khác",
    "Title": "Tiêu đề",
}
arrPolarityTerm = []
driver = GraphDatabase.driver(URI, auth=AUTH)
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

@app.route("/graph", methods=["GET"])
def list():
    global dict
    global arrPolarityTerm
    nodes = []
    edges = []
    with driver.session() as session:
        res_result = session.run(
            "MATCH (start)-[r]->(end) \nRETURN id(start) AS source, id(end) AS target"
        )
        for node in res_result:
            new_edge = {
                "source": "node" + str(node["source"]),
                "target": "node" + str(node["target"]),
            }
            edges.append(new_edge)
        nodes_result = session.run("MATCH (n)\nRETURN id(n) AS id, n.Title AS name")
        for node in nodes_result:
            new_node = {"id": node["id"], "name": node["name"]}
            nodes.append(new_node)
        print(len(nodes))
        print(len(edges))
    return jsonify({"nodes": nodes, "edges": edges})

@app.route("/prompt", methods=["POST"])
def prompt():
    data = request.json
    prompt = data.get("prompt", "")

    if not prompt:
        return jsonify({"error": "Prompt is required"}), 400

    try:
        model = genai.GenerativeModel("gemini-pro")

        response = model.generate_content(prompt)

        print("Response: ", response)
        # return jsonify({"response": response.text})
        return Response(response, mimetype='application/json')
        
    except Exception as e:
        return jsonify({"error": "Request failed"}), 500

def read_file(file_path):
    file_extension = file_path.split('.')[-1].lower()
    
    if file_extension == 'txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    elif file_extension == 'pdf':
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in range(len(reader.pages)):
                text += reader.pages[page].extract_text()
            return text

    elif file_extension == 'docx':
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

    else:
        raise ValueError(f"Unsupported file format: {file_extension}")


model = genai.GenerativeModel("gemini-1.5-flash")

def generate_paper(article_texts):
    prompt = """
    Từ cái bài báo đã cung cấp hãy 
    tổng hợp thành một bài báo mới với nội dung một cách liền mạch, chuyên nghiệp, rõ ràng, chi tiết, bao gồm những cốt lõi của từng bài báo.
    Cuối cùng viết ra một bài báo hoàn chỉnh.
    """

    for index, article in enumerate(article_texts):
        prompt += f"\n\nBài báo {index+1}:\n{article}\n"

    response = model.generate_content(prompt)
    return response.text

@app.route("/paper", methods=["GET"])
def display_paper():
    files = ['378-775-1-SM.pdf', '458-932-1-SM.pdf', 'CVv380S262020085.pdf']
    article_texts = []
    for file in files:
        try:
            content = read_file(file)  # Hàm read_file đã được định nghĩa trước đó
            article_texts.append(content)
        except ValueError as e:
            print(e)
    
    # Tổng hợp nội dung các bài báo bằng Google AI
    if article_texts:
        new_paper = generate_paper(article_texts)
        print("Bài báo đã tổng hợp:\n")
        print(new_paper)

        document = Document()
        document.add_heading('Bài báo tổng hợp', level=1)
        document.add_paragraph(new_paper)
        document.save('result.docx')
        return(new_paper)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)  
